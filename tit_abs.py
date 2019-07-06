import requests
import re
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from requests.exceptions import RequestException
from lxml import etree


def get_html_text(url, headers):
	try:
		r = requests.get(url, headers=headers)
		r.raise_for_status()
		r.encoding = r.apparent_encoding
		return r.text
	except RequestException:
		print('Fail to get html text.')


def parser_lit_links(html):
	pattern = re.compile('.*?class="issue-item_title".*?href="(.*?)">', re.S)
	lit_brok_links = re.findall(pattern, html)
	lit_links = []
	for lit_brok_link in lit_brok_links:
		lit_link = 'https://pubs.acs.org' + lit_brok_link
		lit_links.append(lit_link)
	return lit_links


def get_title(html):
	pattern = re.compile('content="(.*?)"', re.S)
	title_text = re.findall(pattern, html)[1]
	return title_text


def get_abstract_text(html):
	selector = etree.HTML(html)
	abstract_text = selector.xpath('//meta[@name="dc.Description"]/@content')[0]
	return abstract_text


def get_abstract_pic(html):
	try:
		pattern = re.compile('article_abstract-img.*?src="(.*?)"', re.S)
		abstract_gif_broken_url = re.findall(pattern, html)[0]
		abstract_gif_url = "https://pubs.acs.org" + abstract_gif_broken_url
		return abstract_gif_url
	except:
		print("This article has no abstract picture.")
		return False


def str_symbol_out(string):
	"""This fun is aimed to seperate symbol out from a string.

	Another way from Internet:
	import re
	def validateTitle(string):
		rstr = r"[\/\\\:\*\?\"\<\>\|]"
		new_string = re.sub(rstr, "_", title)
		return new_string
	"""

	string_str_lis = list(string)
	number = 0
	for char in string_str_lis:
		if char in ['<', '>', '/', '\\', '|', ':', '"', '*', '?']:
			string_str_lis.pop(number)
		number += 1
	string_normal = ''.join(string_str_lis)
	return string_normal


def write_to_docx(doc, path, title, abstract, abstract_gif_url):
	try:
		doc.add_heading(title)
		doc.add_heading("Abstract:", level=1)

		# doc.styles['Normal'].font.name = "Times New Roman"
		paragraph = doc.add_paragraph(abstract)
		paragraph_format = paragraph.paragraph_format
		paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

		doc.add_heading("Abstract picture:", level=1)
		title_normal = str_symbol_out(title)
		pic_name = path + title_normal + ".gif"
		if abstract_gif_url:
			with open(pic_name, 'wb') as f:
				r = requests.get(abstract_gif_url)
				f.write(r.content)
			doc.add_picture(pic_name, width=Inches(5.0))
		doc.add_page_break()
		doc.save('Title & Abstract.docx')
		print('Success to save info: ' + title[:50] + "...")
	except:
		print("Fail to write info to text.")


def main(kw):
	url = "https://pubs.acs.org/action/doSearch?AllField=" + kw
	headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_13_3) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/65.0.3325.162 Safari/537.36'}
	html = get_html_text(url, headers)
	lit_links = parser_lit_links(html)
	doc = Document()
	doc.save('Title & Abstract.docx')
	path = "pics for Literature " + kw + os.path.sep
	if not os.path.exists(path):
		os.makedirs(path)
	for lit_link in lit_links:
		lit_html = get_html_text(lit_link, headers=headers)
		title = get_title(lit_html)
		abstract = get_abstract_text(lit_html)
		abstract_gif_url = get_abstract_pic(lit_html)
		write_to_docx(doc, path, title, abstract, abstract_gif_url)
	print("Success to generate a report on " + kw + ".")


if __name__ == '__main__':
	keyword = input("Input keyword: ")
	main(keyword)
